import pandas as pd
import io
import requests
import xml.etree.ElementTree as et
import lxml.etree
import json
from tqdm import tqdm
import regex as re
from more_itertools import split_at
from docx import *
from docx.shared import RGBColor
from xml.etree import ElementTree as ET
from collections import Counter
import numpy as np
import math
from itertools import zip_longest

#%%def
def check_if_all_none(list_of_elem):
    """ Check if all elements in list are None """
    result = True
    for elem in list_of_elem:
        if elem is not None:
            return False
    return result

def paragraph_replace_text(paragraph, regex, replace_str):
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    return paragraph

def round_down(x):
    return int(math.floor(x / 100000)) * 100000
#%% modyfikacje pliku .docx
#zamienić: ^t\* na: ^t*^t

#%% nowe podejście

document = Document("C:\\Users\\Cezary\\Downloads\\X. 2  Albania - Peru.docx")
document = Document("C:\\Users\\Cezary\\Downloads\\X.3 Portoryko - łac.docx")

#edycja pliku docx

regex = re.compile("\t\* ")
for paragraph in tqdm(document.paragraphs):
    paragraph_replace_text(paragraph, regex, '\t*\t')
    
#POMYSŁ: wyrzucamy inicjalne gwiazdki i otaczające je tabulatory z tekstu
#NA PRZYSZŁOŚĆ: gwiazdka na końcu oznacza dziedziczenie tłumacza z poprzedniego wersu

#wydobywanie autorów po stylu czcionki
# with open("filename.xml", "w", encoding='utf-8') as f:
#     f.write(test)
#może zliczyć jeszcze rozkład lowercase i uppercase?

total = []
for paragraph in tqdm(document.paragraphs):
    letters = []
    sizes = []
    styles = []
    colours = []
    italic = []
    bold = []
    try:
        tab_value = [e.position for e in paragraph.paragraph_format.tab_stops]#[0]
    except IndexError:
        tab_value = None
    for run in paragraph.runs:
        try:
            sizes.append(run.font.size.pt)
        except AttributeError:
            sizes.append(None)
        colours.append(run.font.color.rgb)
        styles.append(run.font.name)
        letters.append(run.text)
        italic.append(run.italic)
        bold.append(run.bold)
        uppercases_chains = [len(e) for e in re.findall('\p{Lu}+', paragraph.text)]
    total.append([paragraph.text, letters, styles, sizes, colours, italic, bold, uppercases_chains, tab_value])

# test dowolna pozycja z Total: print(total[5])

literatury = []
for ind, el in enumerate(total):
    # el = total[5]
    if RGBColor(0xff, 0x00, 0x00)in el[4]:
        good_indices = [i for i,e in enumerate(el[4]) if e == RGBColor(0xff, 0x00, 0x00)]
        is_italic = [el[5][i] for i in good_indices]
        is_bold = [el[6][i] for i in good_indices]
        if not any(is_italic) and not any(is_bold):
            literatury.append((ind, el[0]))
final_dict = {}
for i, (index, lit) in enumerate(literatury):
    lit = lit.replace("  ", "|").replace(" ", "").replace("|", " ").replace("  ", " ").strip()
    if i in range(len(literatury)-1):
        final_dict[lit] = (index, literatury[i+1][0])
    else:
        final_dict[lit] = (index, ind)
##final_dict czyszczenie białych znaków Literatura: "G R E C J A    w s p ó ł c z e s n a,  ś r e d n i o w i e c z n a,  s t a r o ż y t n a   GR".replace("  ", "|").replace(" ", "").replace("|", " ").replace("  ", " ")


checklists = [['Arial Black'], [9]]

for lit, (i_start, i_end) in final_dict.items():
    # i_start = 260
    # i_end = 780
    lit_list = total[i_start:i_end]
    # paragraph = total[932]
    # paragraph = total[39]
    #7525, 11176, 58
    authors = []
    for index, paragraph in enumerate(lit_list):
        results = []
        # for ind, sublist in enumerate(paragraph[2:4]):
        #     iteration = 0
        #     result = []
        #     while iteration + len(checklists[ind]) <= len(sublist):
        #         result.append(sublist[iteration:iteration+len(checklists[ind])] == checklists[ind])
        #         iteration += 1
        #     if any(result):
        #         result = True
        #         results.append(result)
        #     else: results.append(False)
        # try:
        #     results.append(all([paragraph[3][0] != 10, paragraph[3][1] != 10 if paragraph[1][0] == '\t' else True]))
        # except IndexError:
        #     results.append(False)
        results.append(any(e > 1 for e in paragraph[-2]))
        results.append(check_if_all_none(paragraph[5]) or bool(re.search('Nobel\s+\d+', paragraph[0])) or not any(paragraph[5]))
        results.append(9 in paragraph[3])
        # results.append(all([e == False for e in paragraph[-1]]))
        results.append(any(el in [0,1,2] for el in [i for i, e in enumerate(paragraph[2]) if e == 'Arial Black']))
        try:
            results.append(paragraph[0].strip()[0] != '*')
            results.append(paragraph[0].strip()[-1] != '*')
            results.append(paragraph[0].strip()[0] != '→')
        except IndexError:
            results.append(False)
        if all(results):
            authors.append((index+i_start, paragraph[0]))
    authors = [(e[0], authors[i+1][0] if i in range(len(authors)-1) else i_end, e[-1]) for i, e in enumerate(authors)]
    temp_dict = {}
    for st, end, aut in authors:
        aut = re.sub("\s+", " ", aut.strip())
        temp_dict[aut] = (st, end)
    final_dict[lit] = temp_dict

# spróbować podzielić string przez liczbę elementów w liście z tabulatorami, dzieląc po \t lub wielokrotnej spacji
# spodziewam się otrzymać schodki i chcę dziedziczyć info na wyższym schodku
# start_ind = 291
# end_ind = 443

# start_ind = 454
# end_ind = 568
# # start_ind = 42
# # end_ind = 58

# start_ind = 71
# end_ind = 90

lista_grup = []
for e in tqdm(final_dict):
    for k,v in final_dict[e].items():
        start_ind = v[0]
        end_ind = v[1]
    
        original = [[e[-1][:-1], e[0]] for e in total[start_ind:end_ind]]
        original = [[e[0], e[1], [i for i, el in enumerate(re.finditer('\t', e[1]))]] for e in original]
        original = [[[el for i, el in enumerate(a) if i in c],b,c] for a,b,c in original]
        
        a = [[int(el) for el in e[0]] for e in original]
        if all(not(e) for e in a):
            wartosci = [500000, 600000, 700000, 800000, 900000, 1000000]
            a = [[e] for i, e in enumerate(wartosci) if i in range(len(a))]
        if not(a[0]) and a[1]:
            a[0] = [600000]
            
        b = [e[1] for e in original]
        
        #opowieści o jeźdźcach i pojedynek nie mogą być razem
        #indeksuje się *, trzeba to zbiorczo usunąć
        order = []
        for el1, el2 in zip(a,b):
            # el1 = a[0]
            # el2 = b[0]
            if len(el1) == 1 and el2 != '':
                # order.append((el1[0], el2.replace('*','').strip() if not el2[0] == '.' else el2.replace('*','').replace('.','').strip()))
                order.append((el1[0], el2.strip() if not el2[0] == '.' else el2.replace('.','').strip()))
            elif len(el1) > 1:
                if (list(zip_longest(el1, [e for e in el2.split('\t') if e])) == list(zip(el1, [e for e in el2.split('\t') if e]))) and all(e[-1] for e in list(zip_longest(el1, [e for e in el2.split('\t') if e]))):
                    for subel1, subel2 in zip(el1, [e for e in el2.split('\t') if e]):
                        order.append((subel1, subel2.strip() if not subel2[0] == '.' else subel2.replace('.','').strip()))
                else:
                    for subel1, subel2 in zip(el1, [e for e in el2.split('\t') if not bool(re.search('^\s+$', e)) and e]):
                        order.append((subel1, subel2.strip() if not subel2[0] == '.' else subel2.replace('.','').strip()))
            else:
                try:
                    order.append((min([e for sub in a for e in sub]), el2.strip() if not el2[0] == '.' else el2.replace('.','').strip()))
                except IndexError:
                    pass
        
        order = [(ind, text.replace('♦','').strip()) for ind, text in order]
        #wyznaczanie całostek
        
        test = []
        for i, (ind, tex) in enumerate(order):
            if tex != '' and tex != '*':
                if i == 0 or ((ind < 650000 and not bool(re.search('^\d{4}', tex))) or tex == '——' or bool(re.search('^(\p{Ll} )+.{0,1}$', tex))):
                    test.append([(ind, tex)])
                else: test[-1].append((ind, tex))
        
        #łączenie opisów bibliograficznych
        
        new_test = []
        for el in test:
            # el = test[43]
            if len(set([e[0] for e in el])) > 1:
                max_ind_in_el = max([e[0] for e in el])
            else: max_ind_in_el = 0
            temp_list = []
            for i, (ind, tex) in enumerate(el):
                # i, (ind, tex) = 2, el[2]
                if i != 0 and tex != '——' and ((ind > 1300000 and not bool(re.search('^(\p{Ll} )+.{0,1}$', tex))) or (el[i-1][-1].endswith((';',',')) if el[i-1] != tex else False) or (ind == max_ind_in_el if (not(bool(re.search('s\. {d}', el[i-1][-1]))) and not(bool(re.search('\p{Lu}\.$', el[i-1][-1]))) and len(el[i-1][-1]) > 3) else False) or (el[i-1][-1].endswith(':') if (not(bool(re.search('s\. {d}', tex))) and len(tex) > 3) else False)) and not tex.startswith('='):
                    temp_list[-1].append((ind, tex))
                else: temp_list.append([(ind, tex)])
            new_test.append(temp_list)

        new = []
        for blok in new_test:
            nowy_blok = []
            for lista in blok:
                if len(lista) == 1:
                    nowy_blok.append(lista[0])
                else:
                    ind = lista[0][0]
                    text = ' '.join([e[-1].strip() for e in lista])
                    nowy_blok.append((ind, text))
            new.append(nowy_blok)
            
       #sprawdzić, czy borgesa nie popsuje 
       # czy w literaturze przedmiotowej mam z automatu łączyć nazwiska z rekordami?

# to rozwiązanie dla bloków dłuższych niż 1
       
        for ind, blok in enumerate(new):
            # blok = new[1]
            order = []
            value = 1
            groups = []
            group = 0
            for i, (no, tex) in enumerate(blok):
                # liczba = 0
                # i, (no, tex) = liczba, blok[liczba]
                if i != 0 and no > blok[i-1][0]:
                    value += 1
                    order.append(value)
                    groups.append(group)
                else:
                    value = 1
                    order.append(value)
                    group += 1
                    groups.append(group)
            max_order = max(order)
            
            neww = []
            groups_set = set(groups)
            for group in groups_set:
                # group = 1
                max_in_group = max([o for g, o in zip(groups, order) if g == group])
                if max_in_group < max_order:
                    diff = max_order - max_in_group
                    neww.extend([(g, o+diff) for g, o in zip(groups, order) if g == group])
                else:
                    neww.extend([(g, o) for g, o in zip(groups, order) if g == group])
            
            group_dict = {}
            for group in groups_set:
                # group = 1
                temp_list = [(o[1], t[1]) for t, o in zip(blok, neww) if o[0] == group]
                group_dict.update({group: dict(temp_list)})
             
            for ka, va in group_dict.items():
                # ka = 2
                # va = group_dict[ka]
                if ka > 1:
                    diff = set(group_dict[ka-1].keys()) - set(va.keys())
                    va.update({kal:val for kal,val in group_dict[ka-1].items() if kal in diff})
                
            group_dict = {k:dict(sorted(v.items())) for k,v in group_dict.items()}
            blok = ['|'.join(group_dict[e].values()) for e in group_dict]
            new[ind] = blok
        
        new = [e for sub in new for e in sub]
        # e = 'Argentyna'
        # k = 'Borges'
        temp_dict = {'literatura': e,
                     'pisarz': k,
                     'zapisy bibliograficzne': new[1:]}
        
        lista_grup.append(temp_dict)
                
            
        #Borges wygląda dobrze, teraz zastosować do wszystkich i wprowadzić poziom literatury i autora   
            
for ind, slownik in enumerate(lista_grup):
    # slownik = lista_grup[22]
    byla_gwiazdka = ['nie' for e in range(len(slownik['zapisy bibliograficzne']))]
    for i, zapis in enumerate(slownik['zapisy bibliograficzne']):
        # zapis = slownik['zapisy bibliograficzne'][3]
        try: 
            rok_i_numer = re.findall('\d{4}.+?(?=s\.)', zapis)[0].strip()
        except IndexError:
            pass
        if '~' in zapis:
            lista_grup[ind]['zapisy bibliograficzne'][i] = zapis.replace('~', rok_i_numer)
        try:
            translator = '; ' + re.findall('tł\.[ \|]\p{Lu}.+?\p{Lu}\.(?=,|\|)', zapis)[0].strip().replace('|', ' ')
        except IndexError:
            pass
        if '*' in zapis:
            lista_grup[ind]['zapisy bibliograficzne'][i] = lista_grup[ind]['zapisy bibliograficzne'][i].replace('*', translator)
            byla_gwiazdka[i] = 'tak'
            lista_grup[ind].update({'była gwiazdka': byla_gwiazdka})




df = pd.DataFrame(lista_grup)

df['była gwiazdka'] = df[['była gwiazdka', 'zapisy bibliograficzne']].apply(lambda x: x['była gwiazdka'] if not(isinstance(x['była gwiazdka'], float)) else ['nie' for e in range(len(x['zapisy bibliograficzne']))], axis=1)

df = df.explode(['zapisy bibliograficzne', 'była gwiazdka'])
            
df.to_excel('literatura_na_swiecie_df.xlsx', index=False)            
            
            
            
#%% Praca z df (tabelą)           
            
df = pd.read_excel(r"C:\Users\RivenDell\Documents\$Nauka\HumCyf\PROJEKT\literatura_na_swiecie_df 2022-06-15.xlsx")            
#"typ publikacji"
#"współpracownik - rola"
#"współpracownik - nazwisko"
#"opis bibliograficzny"
#"rok"
#"numer bieżący"
#"numer ciągły"
#"strony"
#"tytuł utworu"
#"dodatkowe informacje" (np. Nobel)

borges_df = df[df["pisarz"]=="BORGES, Jorge Luis 1899–1986"]
def get_typ_publikacji(x):
    try:
        return re.search("(?<!\p{Lu} )(\p{Ll} ){2,}\p{Ll}(?!\p{L})", x).group(0)
    except AttributeError:
        return np.NaN
        
borges_df["typ publikacji"] = borges_df["zapisy bibliograficzne"].apply(lambda x: get_typ_publikacji(x))
def get_wspolpracownik_rola(x):
    if isinstance(x["typ publikacji"], str) and x["zapisy bibliograficzne"].startswith(x["typ publikacji"]) and x["zapisy bibliograficzne"].count("|") > 1:
        return x["zapisy bibliograficzne"].split("|")[1]
    elif isinstance(x["typ publikacji"], str) and x["zapisy bibliograficzne"].startswith(x["typ publikacji"]):
        return x["zapisy bibliograficzne"].split(" ; ")[1].split(" ")[0]
    elif "tł." in x["zapisy bibliograficzne"]:
        return "tł."
        
borges_df["współpracownik - rola"] = borges_df[["zapisy bibliograficzne", "typ publikacji"]].apply(lambda x: get_wspolpracownik_rola(x),axis=1)
def get_wspolpracownik_nazwisko(x):
    if isinstance(x["współpracownik - rola"], str):
        try:
            return re.findall("(?<=tł. |tł.\|)\p{Lu}[\p{L}-]+, \p{Lu}.",x["zapisy bibliograficzne"])[0]
        except IndexError:
            return np.NaN

borges_df[]

borges_df["współpracownik - nazwisko"] = borges_df[["zapisy bibliograficzne", "współpracownik - rola"]].apply(lambda x: get_wspolpracownik_nazwisko(x),axis=1)


y = "w i e r s z e|tł.|Baterowicz, M.|1973 nr 12 (32)  s. 246–7,  Aluzja do śmierci pułkownika Francisco Borgesa ; Ajedrez"
re.findall("(?<=tł. |tł.\|)\p{Lu}[\p{L}-]+, \p{Lu}.",y)[0]

re.match("(?<!\p{Lu} )(\p{Ll} ){2,}\p{Ll}(?!\p{L})","Dziewięć esejów dantejskich ;  tł. Partyka, J.,  Prószyński i S-ka  2007|= Nueve ensayos dantescos r e c.  Mrugalski, M., Ugolino z Argentyny, 2010 nr 1–2 (462–3)  s. 399–418	.").group(0)

### połączenie LnŚ + PBL + BN + VIAF -> połączony dataset

#%%        
        wiersze = new_test[1]
        wiersze_tylko_tekst = []
        for lista in wiersze:
           # lista = wiersze[0]
           if len(lista) > 1:
               wiersze_tylko_tekst.append((lista[0][0], ' '.join([e[-1] for e in lista])))
           else: wiersze_tylko_tekst.append(lista[-1])
           
for i, t in wiersze_tylko_tekst:
    print(i, t)
               
           for indeks, tresc in lista:
               print(tresc)
           print(lista)
        
       
        tekst = 'w i e r s z e|tł.|Baterowicz, M.|1973 nr 12 (32)  s. 246–7,  Aluzja do śmierci pułkownika Francisco Borgesa ; Ajedrez'
        
        tekst.split('|')[3]
       
        new = []
        [new.append(e[0][0][-1]) for e in new_test if len(e) == 1]
       
        
       
        
       
        
       
        
       
        nowe_wiersze = []
        for i, (ind, tex) in enumerate(wiersze):
            if ind > 1300000 or wiersze[i-1][-1].endswith((';',',')):
                nowe_wiersze[-1].append((ind, tex))
            else: nowe_wiersze.append([(ind, tex)])
                
        rex = 'W osiemdziesiąt światów dookoła dnia,  Czyt. 1976'
        tex = 'La vuelta al dia en ochenta mundos'
        ind = 895350
        
        tex != '——' and (ind > 1300000 or rex.endswith((';',',')) or (ind == max_ind_in_el if not(bool(re.search('s\. {d}', rex))) else False) or (rex.endswith(':') if not(bool(re.search('s\. {d}', tex))) else False)) and not tex.startswith('=')
        
        
        
        
                         
                         '~' not in tex and (not(bool(re.search('\d{4}', tex))) and not(bool(re.search('\d{4}', rex)))) else False)
        
        
        bool(re.search('s\. {d}', tex)) and not(bool(re.search('s\. {d}', rex)))
        
        rex = 'Chorób nie ma, są tylko chorzy,  ~ s. 276–7'
        tex = 'Chorób nie ma, są tylko chorzy,  ~ s. 276–7'
        
        
        
        tex != '——' and (ind > 1300000 or el[i-1][-1].endswith((';',',')) or ind == max_ind_in_el if '~' not in tex and (not(bool(re.search('\d{4}', tex))) and not(bool(re.search('\d{4}', rex)))) else False)
                
                
                
    # opracować warunki łączenia treści w całostki bibliograficzne
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
        # for e, f in zip(test, text):
        #     if e[-1].strip() != f:
        #         print(f'{e} | {f}')
        
        # text = [e for e in '\n'.join([e[0] for e in total[start_ind:end_ind]]).split('\t') if e]
        # text = [e for e in text if not(re.match('\s', e[0]))]
        # test = [e[0] for e in original if e[0]]
        # test = [int(e) for sub in test for e in sub]
                    
        text = [e[-1] for e in order]
        order = [e[0] for e in order]
        
        groups = []
        counter = 0
        for ind, number in enumerate(order):
            # if ind == 0 or number > order[ind-1] or number > 1300000:
            if ind == 0 or number > order[ind-1]:
                groups.append(counter)
            else:
                counter += 1
                groups.append(counter)
        
        # groups_set = set(groups)           
        # # groups_dict = {}
        # indices_out = []
        # for s in groups_set:
        #     # s=0
        #     group_indices = [i for i, e in enumerate(groups) if e == s]
        #     single_group = [e for i, e in enumerate(order) if i in group_indices]
        #     ind_list = [e for e,f in zip(group_indices, single_group) if f > 1300000]
        #     if ind_list:
        #         ind_list.insert(0,ind_list[0]-1)
        #         merged_str = ' '.join([e for i, e in enumerate(text) if i in ind_list]) 
        #         text[ind_list[0]] = merged_str
        #         indices_out.extend(ind_list[1:])
        # indices_out.extend([i for i, e in enumerate(text) if e in ['', '*']])
                        
        # text = [e for i,e in enumerate(text) if i not in indices_out]  
        # groups = [e for i,e in enumerate(groups) if i not in indices_out]
        # order = [round_down(e) for i,e in enumerate(order) if i not in indices_out]
        
        order = [round_down(e) for i,e in enumerate(order)]
        order_indices = [(e,i) for i,e in enumerate(sorted(list(set(order))),1)]
        order_indices = dict(zip([e[0] for e in order_indices], [e[1] for e in order_indices]))
        order_indices = [order_indices[e] for e in order]
        
        # tu trzeba by połączyć to, co stanowi całość
        
        groups_set = set(groups)
        new_order = []
        for group in groups_set:
            group = 4
            print([o for o, g in zip(order_indices, groups) if g == group])
            length = len([o for o, g in zip(order_indices, groups) if g == group])
            min_value = [o for o, g in zip(order_indices, groups) if g == group][0]
            max_value = [o for o, g in zip(order_indices, groups) if g == group][-1]
            old_set = set([o for o, g in zip(order_indices, groups) if g == group])
            
            new_set = set(list(range(min_value, max_value+1)[:length]))
            if dict(zip(old_set-new_set, new_set-old_set)):
                difference = dict(zip(old_set-new_set, new_set-old_set))
            update = [difference[e] if e in difference else e for e in list(range(min_value, max_value+1)[:length])]
            
            new_order.extend(update)
            print(new_order)
        
        
        total_info = list(zip(new_order, text, groups))
        
        biblio_dict = {}
        for ind, string, group in total_info:
            if group not in biblio_dict:
                biblio_dict[group] = {ind: string}
            elif ind in biblio_dict[group]:
                ind -= 1
                biblio_dict[group].update({ind: string})
            else:
                biblio_dict[group].update({ind: string})
        
        df = pd.DataFrame.from_dict(biblio_dict, orient='index').sort_index()
        
        lista_grup.append(biblio_dict)

test = [list(e.values()) for e in lista_grup]        
test = [e for sub in test for e in sub]
df = pd.DataFrame(test)
columns = sorted(df.columns.values)
df = df.reindex(columns=columns)
        # teksty = []
        # for group in groups_set:
        #     test = '|'.join([te for gr, te in zip(groups, text) if gr == group])
        #     teksty.append(test)
        
        
        
        
        
        
        
        
        
        
        
        
        #tu trzeba uważać, bo jeśli w rekordzie jest 9 i 10, to się nadpisze, ale u Borgesa nie ma, więc u kogo może być?
        order = [e if e != 1000000 else 900000 for e in order]
        
        set_order = sorted(list(set(order)))
        if len(set_order) == 1:
            wartosci = [500000, 600000, 700000, 800000, 900000, 1000000]
            order = [e for i, e in enumerate(wartosci) if i in range(len(order))]
            set_order = sorted(list(set(order)))
        
        # tu sprawdzam, czy w grupie frazy się na siebie nie nakładają
        # for g in groups_set:
        #     if any(el[-1] > 1 for el in Counter([f for e,f in zip(groups, order) if e == g]).most_common(1)):
        #         print(g)
            
        total_info = list(zip(order, text, groups))
        
        
        biblio_dict = {}
        for ind, string, group in total_info:
            if group not in biblio_dict:
                biblio_dict[group] = {ind: string}
            elif ind in biblio_dict[group]:
                ind -= 100000
                biblio_dict[group].update({ind: string})
            else:
                biblio_dict[group].update({ind: string})       
        
        # test = biblio_dict[41]
        # list(test.items())
        
        #edycje pól może na później
        # v_to_change = [[el for el in e.values()][-1] for e in biblio_dict.values()]
        # v_changed = v_to_change[:]
        # while any(e[0] == '~' for e in v_changed):
        #     v_changed = [e if e[0] != '~' else f'{v_changed[i-1].split("s.")[0]}{e[2:]}' for i, e in enumerate(v_changed)]
        # change_dict = dict(zip(v_to_change, v_changed))
        
        # biblio_dict = {k:{ke:change_dict[va] if va in change_dict else va for ke,va in v.items()} for k,v in biblio_dict.items()}
        
        
        df = pd.DataFrame.from_dict(biblio_dict, orient= 'index').sort_index()
        columns = sorted(df.columns.values)
        df = df.reindex(columns=columns)
        # df['validate'] = df[4].apply(lambda x: 'ok' if x[:4].isnumeric() else 'not ok')
        # df = df.replace(r'\*', np.nan, regex=True)
        try:
            for i, row in df.iterrows():
                if not(isinstance(row[600000],float)) and isinstance(row[700000],float) and not(isinstance(row[800000],float)):
                    df.at[i, 600000] = f"{row[600000]} {row[800000]}"
                    df.at[i, 800000] = np.nan
        except KeyError:
            pass
                        
        try:
            for i, row in df.iterrows():
                # row = df.loc[80]
                if not(isinstance(row[600000],float)) and isinstance(row[700000],float) and isinstance(row[800000],float) and not(isinstance(row[900000],float)) and row[900000][:5] != 'r e c':
                    df.at[i, 600000] = f"{row[600000]} {row[900000]}"
                    df.at[i, 900000] = np.nan
        except KeyError:
            pass
        
        if df.shape[0] > 1:
            df[500000] = df[500000].fillna(method='ffill')
            df[600000] = df[600000].fillna(method='ffill')
        
        try:
            grouped = df.groupby([500000, 600000], as_index=False).apply(lambda group: group.ffill())
        except KeyError:
            grouped = df.groupby([500000], as_index=False).apply(lambda group: group.ffill())
        
        grupy_dict = {5: 1,
                     2: 2,
                     'reszta': 3}
        
        grouped['grupowanie'] = grouped.apply(lambda x: grupy_dict[len(x.dropna())] if len(x.dropna()) in grupy_dict else grupy_dict['reszta'], axis=1)
        
        # lista_grup.append(Counter(grouped['grupowanie']))
        lista_grup.append(grouped)

df = pd.concat([e for e in lista_grup])

sum(lista_grup, Counter())

for i, row in grouped.iterrows():
    # i = 29
    # row = grouped.loc[i]
    for ind, pole in row.iteritems():
        # ind = 600000
        # print(ind)
        # ind = 900000
        if isinstance(pole, str):
            if '~' in pole:
                # pole = grouped.loc[i][ind]
                try:
                    replacement = re.findall('\d{4}.+?\)', grouped.loc[i-1][ind])[0]
                except IndexError:
                    pass
                except TypeError:
                    try:
                        replacement = re.findall('\d{4}.+?\)', grouped.loc[i][600000])[0]
                    except (IndexError, TypeError):
                        replacement = re.findall('\d{4}.+?\)', grouped.loc[i-1][600000])[0]
                grouped.loc[i,ind] = grouped.loc[i][ind].replace('~', replacement)
for i, row in grouped.iterrows():
    for ind, pole in row.iteritems():
        if isinstance(pole, str):
            if '*' in pole:
                # translator = [e for e in grouped.loc[28] if isinstance(e,str)]
                translator = [e for e in grouped.loc[i-1] if isinstance(e,str)]
                try:
                    translator = [re.findall('tł\. \p{Lu}.+?\p{Lu}\.(?=,)', e)[0] for e in translator if re.findall('tł\. \p{Lu}.+?\p{Lu}\.(?=,)', e)][0]
                except IndexError:
                    translator = 'tł. ' + [re.findall('\p{Lu}.+?\p{Lu}\.(?=,)', e)[0] for e in translator if re.findall('\p{Lu}.+?\p{Lu}\.(?=,)', e)][0]
                grouped.loc[i,ind] = grouped.loc[i][ind].replace('*', translator)
                


# df = df.groupby([600000], as_index=False).apply(lambda group: group.ffill())
grouped.to_excel('borges.xlsx', index=False)


#dalsze kroki
# - wykorzystać podział na grupy i dla każdej z grup zaproponować parsowanie materiału (przy pomocy struktury, np. wiersze, lub przy pomocy regexów)



















    
#jeśli grupa ma 6 elementów, to w text trzeba 2 ostatnie scalić, a w test i groups ostatni element usunąć

indices_out = []
for s in groups_set:
    single_group = [e for e in groups if e == s]
    print(single_group)
    
    
    group_len = len(single_group)
    if group_len == 6:
        indices = [i for i,e in enumerate(groups) if e == s][-2:]
        text[indices[0]] = ' '.join([e.strip() for i,e in enumerate(text) if i in indices])
        indices_out.append(indices[-1])

#usunąć elementy obiektów, które mają indeksy w zmiennej indices_out + jak połączyć rzeczy na tym samym wcięciu lub dwóch ostatnich wcięciach?

del [e for i, e in enumerate(test) if e in indices_out]       
del [e for i, e in enumerate(text) if e in indices_out]
del [e for i, e in enumerate(groups) if e in indices_out]
        
    
temp_dict = dict(Counter(groups))
new_dict = {}
max_length = Counter(groups).most_common(1)[-1][-1]
for k,v in temp_dict.items():
    if v == max_length:
        new_dict[k] = list(range(v))
    else:
        diff = max_length - v
        new_dict[k] = list(range(diff,v+diff))

order = [e for sub in new_dict.values() for e in sub]

for ind, (o, t) in enumerate(zip(order, text)):
    if o == 4 and order[ind+1] == 5:
        text[ind] = ' '.join([text[ind],text[ind+1]])
        
ttt = [(order[ind],text[ind],groups[ind]) for ind,(g,o,t) in enumerate(zip(groups,order,text)) if o != 5]
    
        
    
biblio_dict = {}
for ind, string, group in ttt:
    if group not in biblio_dict:
        biblio_dict[group] = {ind: string}
    else:
        biblio_dict[group].update({ind: string})       

df = pd.DataFrame.from_dict(biblio_dict, orient= 'index').sort_index()
columns = sorted(df.columns.values)
df = df.reindex(columns=columns)
df.to_excel('test.xlsx', index=False)



test = [e[-1][:-1] for e in total[43:58] if e[-1][:-1]]
test = [e for sub in test for e in sub]
del test[-12]


test = [(e[0][:-1], e[-1]) for e in test]
test = [e for e in test if e[0]]
[e[0] for e in test]

# [(e[-1], e[0]) for e in total[7419:7441]]

lista = []
for ind, (t1, t2) in enumerate(test):
    if t1 >= test[ind-1][0]:
        lista[-1].append(t2)
    else:
        lista.append([t2])

 mrk_list = []
    for row in marc_list:
        if row.startswith('=LDR'):
            mrk_list.append([row])
        else:
            if row:
                mrk_list[-1].append(row)
## test czyszczenie białych znaków Autor re.sub("\s+", " ", "	BASHA,  Eqrem   1948–                     ".strip())
    
info_o_publikacjach = []
errors = []

for lit in tqdm(final_dict):
    for aut, (i_start, i_end) in final_dict[lit].items():
        # i_start = 28
        # i_end = 34
        
        aut_list = total[i_start:i_end]
        
        zakres_autora = [e[0] for e in aut_list]
        zakres_autora = [e.strip() for e in zakres_autora]
        try:
            zakres_autora = [f'{e} {zakres_autora[i+1]}' if e[-1] == ';' else e for i,e in enumerate(zakres_autora) if i != [ind for ind,el in enumerate(zakres_autora) if el[-1] == ';'][0]+1]
        except IndexError: pass
        
        liczba_publikacji = len([e for e in zakres_autora if e.strip() == '.'])
        indeksy_publikacji = [e[0] for e in enumerate(zakres_autora) if e[1].strip() == '.']
        dlugosc_listy_osoby = [e[0] for e in enumerate(zakres_autora)]
        zakresy_publikacji = [e for e in list(split_at(dlugosc_listy_osoby[2:], lambda x: x in indeksy_publikacji)) if e]
        try:
            if zakres_autora[1] == 'w i e r s z e':
                for publikacja in zakresy_publikacji:
                    # publikacja = zakresy_publikacji[2]
                    autor = zakres_autora[0]
                    typ_pub = zakres_autora[1]
                    try:
                        wspolpracownik_rola = re.findall("^.+?\t", zakres_autora[publikacja[0]])[0].strip()
                        do_usuniecia = re.findall("^.+?\t", zakres_autora[publikacja[0]])[0]
                    except IndexError:
                        pass
                    wspolpracownik_nazwisko = re.sub("^(.+?)(\p{Lu}.+)", r"\2", zakres_autora[publikacja[0]])
                    wspolpracownik_nazwisko = zakres_autora[publikacja[0]].replace(do_usuniecia, "")
                    for opis in publikacja[1:]:
                        opis_bibliograficzny = zakres_autora[opis].split(", ")[0].strip()
                        if opis_bibliograficzny[0] != "~":
                            rok_i_numer = opis_bibliograficzny.split("s.")[0].strip()
                        else:
                            opis_bibliograficzny = opis_bibliograficzny.replace("~", rok_i_numer)
                    
                        tytuly = zakres_autora[opis].split(", ", maxsplit=1)[1].strip()
                        temp_dict = {"autor": autor,
                                     "typ publikacji": typ_pub,
                                     "współpracownik - rola": wspolpracownik_rola,
                                     "współpracownik - nazwisko": wspolpracownik_nazwisko,
                                     "opis bibliograficzny": opis_bibliograficzny,
                                     "tytuły utworów": tytuly}
                        info_o_publikacjach.append(temp_dict)
            else: errors.append((i_start, i_end, zakres_autora))
        except IndexError: errors.append((i_start, i_end, zakres_autora))
        
#tabela z danymi
df = pd.DataFrame(info_o_publikacjach)
df.to_excel("tabela.xlsx", index = False)    




##do uporządkowania: 1)publikacje uznane za autora na liście autorów - brak autorow PEPETELA 214 2)parsowanie całości 
###inne: 1)oznaczenie kolorem niebieskim (kolor inny niz czarny lub czerwony)

#TUTAJ
    
literatura = test['A L B A N I A   AL         ']

info_o_publikacjach = []
errors = []
for key, value in literatura.items():
    # key = '\tBASHA,  Eqrem   1948–                     '
    # value = literatura[key]
    zakres_autora = [e[0] for e in total[value[0]:value[-1]]]
    zakres_autora = [e.strip() for e in zakres_autora]
    
    liczba_publikacji = len([e for e in zakres_autora if e.strip() == '.'])
    indeksy_publikacji = [e[0] for e in enumerate(zakres_autora) if e[1].strip() == '.']
    dlugosc_listy_osoby = [e[0] for e in enumerate(zakres_autora)]
    zakresy_publikacji = [e for e in list(split_at(dlugosc_listy_osoby[2:], lambda x: x in indeksy_publikacji)) if e]
    try:
        if zakres_autora[1] == 'w i e r s z e':
            for publikacja in zakresy_publikacji:
                # publikacja = zakresy_publikacji[2]
                autor = zakres_autora[0]
                typ_pub = zakres_autora[1]
                try:
                    wspolpracownik_rola = re.findall("^.+?\t", zakres_autora[publikacja[0]])[0].strip()
                    do_usuniecia = re.findall("^.+?\t", zakres_autora[publikacja[0]])[0]
                except IndexError:
                    pass
                wspolpracownik_nazwisko = re.sub("^(.+?)(\p{Lu}.+)", r"\2", zakres_autora[publikacja[0]])
                wspolpracownik_nazwisko = zakres_autora[publikacja[0]].replace(do_usuniecia, "")
                for opis in publikacja[1:]:
                    print(opis)
                    opis_bibliograficzny = zakres_autora[opis].split(", ")[0].strip()
                    if opis_bibliograficzny[0] != "~":
                        rok_i_numer = opis_bibliograficzny.split("s.")[0].strip()
                    else:
                        opis_bibliograficzny = opis_bibliograficzny.replace("~", rok_i_numer)
                
                    tytuly = zakres_autora[opis].split(", ", maxsplit=1)[1].strip()
                    temp_dict = {"autor": autor,
                                 "typ publikacji": typ_pub,
                                 "współpracownik - rola": wspolpracownik_rola,
                                 "współpracownik - nazwisko": wspolpracownik_nazwisko,
                                 "opis bibliograficzny": opis_bibliograficzny,
                                 "tytuły utworów": tytuly}
                    info_o_publikacjach.append(temp_dict)
        else: errors.append(zakres_autora) 
    except IndexError: errors.append(zakres_autora)

    
info_o_publikacjach = []
errors = []
for author in albania_lista[1:]:
    # author = albania_lista[1]
    author = [e.strip() for e in author.split("\n") if e!= ""]
    liczba_publikacji = len([e for e in author if e == '.'])
    indeksy_publikacji = [e[0] for e in enumerate(author) if e[1] == '.']
    dlugosc_listy_osoby = [e[0] for e in enumerate(author)]
    zakresy_publikacji = [e for e in list(split_at(dlugosc_listy_osoby[2:], lambda x: x in indeksy_publikacji)) if e]
    if author[1] == 'w i e r s z e':
        for publikacja in zakresy_publikacji:
            # publikacja = zakresy_publikacji[2]
            autor = author[0]
            typ_pub = author[1]
            try:
                wspolpracownik_rola = re.findall("^.+?\t", author[publikacja[0]])[0].strip()
                do_usuniecia = re.findall("^.+?\t", author[publikacja[0]])[0]
            except IndexError:
                pass
            wspolpracownik_nazwisko = re.sub("^(.+?)(\p{Lu}.+)", r"\2", author[publikacja[0]])
            wspolpracownik_nazwisko = author[publikacja[0]].replace(do_usuniecia, "")
            for opis in publikacja[1:]:
                print(opis)
                opis_bibliograficzny = author[opis].split(", ")[0].strip()
                if opis_bibliograficzny[0] != "~":
                    rok_i_numer = opis_bibliograficzny.split("s.")[0].strip()
                else:
                    opis_bibliograficzny = opis_bibliograficzny.replace("~", rok_i_numer)
            
                tytuly = author[opis].split(", ", maxsplit=1)[1].strip()
                temp_dict = {"autor": autor,
                             "typ publikacji": typ_pub,
                             "współpracownik - rola": wspolpracownik_rola,
                             "współpracownik - nazwisko": wspolpracownik_nazwisko,
                             "opis bibliograficzny": opis_bibliograficzny,
                             "tytuły utworów": tytuly}
                info_o_publikacjach.append(temp_dict)
    else: errors.append(author 

#tabela z danymi
df = pd.DataFrame(info_o_publikacjach)
df.to_excel("tabela.xlsx", index = False)

#%% textract = praca z plikami word



# Create list of paths to .doc files
#w ścieżkach podwójne backslashe lub pojedyncze slashe
file_path = "C:\\Users\\RivenDell\\Documents\\$Nauka\\HumCyf\\PROJEKT\\pliki od p. Bednarz\\lns_test.docx"


text = textract.process(file_path)
test = text.decode("utf-8")

test2 = re.split('\n\t(?=\p{Lu} )+', test)

albania = test2[1]
#regex poniżej do poprawy
# albania_lista = re.split('\n\t{2}(?=\p{Lu}{2,}.+\d+\–)', albania)
albania_lista = re.split('\n\t{2}(?=\p{Lu}{2,},  \p{Lu})', albania)
albania_lista = [re.sub("(;)(\n+\t+)",r"\1 ",e) for e in albania_lista]
info_o_publikacjach = []
errors = []
for author in albania_lista[1:]:
    # author = albania_lista[1]
    author = [e.strip() for e in author.split("\n") if e!= ""]
    liczba_publikacji = len([e for e in author if e == '.'])
    indeksy_publikacji = [e[0] for e in enumerate(author) if e[1] == '.']
    dlugosc_listy_osoby = [e[0] for e in enumerate(author)]
    zakresy_publikacji = [e for e in list(split_at(dlugosc_listy_osoby[2:], lambda x: x in indeksy_publikacji)) if e]
    if author[1] == 'w i e r s z e':
        for publikacja in zakresy_publikacji:
            # publikacja = zakresy_publikacji[2]
            autor = author[0]
            typ_pub = author[1]
            try:
                wspolpracownik_rola = re.findall("^.+?\t", author[publikacja[0]])[0].strip()
                do_usuniecia = re.findall("^.+?\t", author[publikacja[0]])[0]
            except IndexError:
                pass
            wspolpracownik_nazwisko = re.sub("^(.+?)(\p{Lu}.+)", r"\2", author[publikacja[0]])
            wspolpracownik_nazwisko = author[publikacja[0]].replace(do_usuniecia, "")
            for opis in publikacja[1:]:
                print(opis)
                opis_bibliograficzny = author[opis].split(", ")[0].strip()
                if opis_bibliograficzny[0] != "~":
                    rok_i_numer = opis_bibliograficzny.split("s.")[0].strip()
                else:
                    opis_bibliograficzny = opis_bibliograficzny.replace("~", rok_i_numer)
            
                tytuly = author[opis].split(", ", maxsplit=1)[1].strip()
                temp_dict = {"autor": autor,
                             "typ publikacji": typ_pub,
                             "współpracownik - rola": wspolpracownik_rola,
                             "współpracownik - nazwisko": wspolpracownik_nazwisko,
                             "opis bibliograficzny": opis_bibliograficzny,
                             "tytuły utworów": tytuly}
                info_o_publikacjach.append(temp_dict)
    else: errors.append(author) #w errors przypadki: artykuł w nrze LnŚ oraz publikacja (książka) z adnotacją o nagrodzie
    #kolejny typ: recenzja
    #kolejne: uruchomić pętlę dla wszystkich autorów (podział publikacji)
basha = albania_lista[1]

basha = albania_lista[2]

#rozdzielić autora na części
#kod poniżej usuwał pojedynczą kropkę w nowym wierszu
# test_autor = [e.strip() for e in basha.split("\n") if e!= "" and e.strip() != "."]
test_autor = [e.strip() for e in basha.split("\n") if e!= ""]

autor = test_autor[0]

# do zrobienia: lata_zycia
typ_pub = test_autor[1]
#pętla - wiele utworów / ile kropek, tyle utworów
liczba_publikacji = len([e for e in test_autor if e == '.'])
# co jest kropką - która pozycja na liście
indeksy_publikacji = [e[0] for e in enumerate(test_autor) if e[1] == '.']
#w miarę oczekiwany wynik: [(2,3),(5,6)]
#jak działają tuple: indeks[0] + treść[1]
# indeksy_publikacji = [e for e in enumerate(test_autor)]
#ile elementów jest w jednym autorze [-1] zwróci ostatnią wartość
dlugosc_listy_osoby = [e[0] for e in enumerate(test_autor)]

zakresy_publikacji = [e for e in list(split_at(dlugosc_listy_osoby[2:], lambda x: x in indeksy_publikacji)) if e]

info_o_publikacjach = []
for publikacja in zakresy_publikacji:   
    autor = test_autor[0]
    typ_pub = test_autor[1]
    try:
        wspolpracownik_rola = re.findall("^.+?\t", test_autor[publikacja[0]])[0].strip()
        do_usuniecia = re.findall("^.+?\t", test_autor[publikacja[0]])[0]
    except IndexError:
        pass
    wspolpracownik_nazwisko = re.sub("^(.+?)(\p{Lu}.+)", r"\2", test_autor[publikacja[0]])
    wspolpracownik_nazwisko = test_autor[publikacja[0]].replace(do_usuniecia, "") 
    opis_bibliograficzny = test_autor[publikacja[1]].split(", ")[0].strip()
    if opis_bibliograficzny[0] != "~":
        rok_i_numer = opis_bibliograficzny.split("s.")[0].strip()
    else:
        opis_bibliograficzny = opis_bibliograficzny.replace("~", rok_i_numer)

    tytuly = test_autor[publikacja[1]].split(", ")[1].strip()
    temp_dict = {"autor": autor,
                 "typ publikacji": typ_pub,
                 "współpracownik - rola": wspolpracownik_rola,
                 "współpracownik - nazwisko": wspolpracownik_nazwisko,
                 "opis bibliograficzny": opis_bibliograficzny,
                 "tytuły utworów": tytuly}
    info_o_publikacjach.append(temp_dict)
# rok = 
# numer
# numer_ciagly
# strony

#można dodać podział utworów: .split(" ; ")

#słownik dla jednej osoby
one_person_dict = {"autor": autor,
                   "typ publikacji": typ_pub,
                   "współpracownik - rola": wspolpracownik_rola,
                   "współpracownik - nazwisko": wspolpracownik_nazwisko,
                   "opis bibliograficzny": opis_bibliograficzny,
                   "tytuły utworów": tytuly}

#tabela dla jednej osoby
one_person_df = pd.DataFrame([one_person_dict])
one_person_df.to_excel("tabela.xlsx", index = False)

#tabela dla wielu publikacji jednej osoby
publications_df = pd.DataFrame(info_o_publikacjach)

#tabela dla wielu osób
people_publications_df = pd.DataFrame(info_o_publikacjach)




#%% praca na pliku PDF
# with open("lns1_1971-2014.pdf", 'wb') as fd:
#             fd.write(r.content)
with pdfplumber.open("lns1_1971-2014.pdf") as pdf:
    pl_txt = ''
    for page in tqdm(pdf.pages):
        pl_txt += '\n' + page.extract_text()
txt = io.open("lns1_1971-2014.txt", 'wt', encoding='UTF-8')
txt.write(pl_txt)
txt.close()


test = io.open('lns1_1971-2014.txt', encoding='utf8').readlines()


lista = []
for row in test:
    try:
        if re.findall('^\p{Lu} \p{Lu} \p{Lu}', row):
            lista.append([row])
        else:
            if row:
                lista[-1].append(row)
    except IndexError:
        pass

slownik = {}
for row in test:
    try:
        if re.findall('^\p{Lu} \p{Lu} \p{Lu}', row):
            nazwa_literatury = re.findall('^\p{Lu} \p{Lu} \p{Lu}.+', row)[0].strip()
            slownik[nazwa_literatury] = [row]
        else:
            if row:
                slownik[nazwa_literatury].append(row)
    except (KeyError, NameError):
        pass

argentyna = slownik['A R G E N T Y N A   AR']
argentyna_lista_list = []
for index, lista in enumerate(argentyna):
    try:
        if re.findall('^ {0,10}\.+', lista):
            argentyna_lista_list.append([lista])
        else:
            argentyna_lista_list[-1].append(lista)
    except IndexError:
        pass
        

# for k,v in slownik.items():
#     for element in v:
#         try:
#             if re.findall('^ {0,10}\.+', element):
#                 slownik[k].append([element])
#             else:
#                 slownik[k][-1].append(element)
#         except (TypeError, AttributeError):
#             pass













ttt = test[:20000]

lista = re.split('(^\p{Lu} \p{Lu})', ttt)
